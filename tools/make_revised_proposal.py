from docx import Document
from docx.shared import Pt

source = "file/开题报告_修改定稿版.docx"
target = "file/开题报告_按反馈修改版.docx"

doc = Document(source)

replacements = {
    20: "本项目以“Julia”AI陪伴终端为研究载体，基于微雪 ESP32-S3 1.85 寸 LCD 开发板，研究资源受限嵌入式设备上的端云协同智能交互方法。系统以语音、惯性运动、时间与设备状态等信息作为环境上下文，由端侧完成唤醒检测、状态管理、界面表达、长期记忆与设备控制，云端完成语音识别、大语言模型推理和语音合成。研究重点包括：构建“感知—决策—表达”闭环架构；设计包含 6 个主状态和 20 个子状态的层次状态机，实现被动问答、主动关怀与用户拒绝后的节制交互；研究 Function Call 与确定性设备控制之间的安全映射，并通过可重复实验评价系统的实时性、可靠性和交互体验。",
    32: "本项目旨在设计并实现一套面向情感陪伴的端云协同多模态智能交互终端（代号“Julia”），实现以下目标：（1）构建可测量、可降级的中文语音交互链路，在稳定网络下统计端到端时延、识别成功率与失败恢复时间；（2）设计由情感线索和环境上下文驱动的主动交互状态机，实现从被动响应到适度主动陪伴的转变；（3）实现大模型 Function Call 与设备控制抽象层的融合，并对控制结果进行校验和反馈；（4）在 ESP32-S3 平台完成系统部署，通过连续运行、弱网、存储和功耗实验验证系统可靠性。",
    40: "（1）边缘端资源约束下的多模态实时融合问题。ESP32-S3 的片上 SRAM、PSRAM 带宽和处理能力有限，如何在不影响语音前端的前提下协调音频采集、网络传输、图形渲染、状态推理和存储访问，是本研究的核心问题。",
    41: "（2）主动陪伴策略的可解释决策与边界控制问题。如何融合语音活动、运动、时间、历史交互和用户反馈，在主动关怀与避免打扰之间取得平衡，并使状态转换可解释、可复现，是交互策略设计的关键问题。",
    42: "（3）端云协同的流式交互可靠性问题。如何在 DNS 失败、网络抖动、云端超时和服务不可用时进行有界等待、重试退避和离线降级，并通过实验量化恢复时间和用户可感知影响，是系统落地必须解决的问题。",
    45: "系统采用端云协同架构。ESP32-S3 端负责传感数据采集、本地语音前端、层次状态机、长期记忆、界面渲染、功耗管理和设备控制；云端负责 ASR、LLM 与 TTS。SD 卡和 NVS 用于对话摘要、画像、状态与离线语音缓存，网络层提供 TLS、DNS 失败回退、超时重试、信号质量检测和弱网降级。系统总体架构如图 1 所示。",
    57: "（1）功能验证：测试唤醒、录音、ASR、LLM、TTS、多轮对话、长期记忆、情绪状态联动和设备控制。在统一语料、固定网络和相同音量条件下，分别统计唤醒成功率、ASR 字错误率、端到端时延中位数/P95、失败恢复时间及控制成功率；每项指标至少重复 30 次，避免以单次演示结果替代量化评测。",
    65: "（1）学术成果：完成硕士学位论文 1 篇；围绕端云协同交互、主动陪伴状态管理或嵌入式系统可靠性形成论文、技术报告或专利成果，具体投稿层级根据实验数据完整性和创新性确定。",
}

for index, text in replacements.items():
    if index >= len(doc.paragraphs):
        raise RuntimeError(f"paragraph index out of range: {index}")
    doc.paragraphs[index].text = text

# Keep the report date and schedule consistent with the current development stage.
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if "2026年7月" in cell.text:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = "宋体"
                        run.font.size = Pt(10.5)

doc.core_properties.title = "面向情感陪伴的端云协同多模态智能交互终端设计与实现（按反馈修改版）"
doc.core_properties.comments = "根据指导教师反馈，补充研究边界、可解释状态策略、弱网可靠性与量化实验方案。"
doc.save(target)
print(target)
