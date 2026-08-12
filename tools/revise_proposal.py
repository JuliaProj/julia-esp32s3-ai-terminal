from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "file" / "开题报告.docx"
BACKUP = ROOT / "file" / "开题报告_原稿备份.docx"
OUTPUT = ROOT / "file" / "开题报告_修改定稿版.docx"
DIAGRAM = ROOT / "file" / "Julia系统总体架构图.png"


def font(size, bold=False):
    candidates = [
        Path(r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def rounded_box(draw, xy, fill, outline, title, lines):
    draw.rounded_rectangle(xy, radius=16, fill=fill, outline=outline, width=3)
    x1, y1, x2, _ = xy
    draw.text(((x1 + x2) / 2, y1 + 18), title, font=font(28, True), fill="#172033", anchor="ma")
    y = y1 + 62
    for line in lines:
        draw.text(((x1 + x2) / 2, y), line, font=font(21), fill="#263247", anchor="ma")
        y += 34


def arrow(draw, start, end, label=None):
    draw.line([start, end], fill="#516078", width=5)
    x2, y2 = end
    draw.polygon([(x2, y2), (x2 - 12, y2 - 8), (x2 - 12, y2 + 8)], fill="#516078")
    if label:
        draw.text(((start[0] + end[0]) / 2, start[1] - 13), label,
                  font=font(18), fill="#516078", anchor="ms")


def make_diagram():
    image = Image.new("RGB", (1800, 1120), "white")
    draw = ImageDraw.Draw(image)
    draw.text((900, 38), "Julia 端云协同多模态智能交互系统总体架构",
              font=font(38, True), fill="#172033", anchor="ma")

    rounded_box(draw, (80, 150, 500, 440), "#EAF4F1", "#27856F", "感知与执行层", [
        "I2S 数字麦克风 / 功放", "1.85 英寸 LCD / RGB565", "IMU、RTC、电池 ADC", "Wi-Fi / BLE / 家居设备接口"
    ])
    rounded_box(draw, (690, 110, 1110, 480), "#EDF2FA", "#3B6DA8", "ESP32-S3 边缘智能层", [
        "FreeRTOS 多任务调度", "WakeNet + VAD 本地语音前端", "6 主状态 / 20 子状态 HFSM", "长期记忆、用户画像与主动关怀", "LVGL 状态动画与低功耗管理"
    ])
    rounded_box(draw, (1300, 150, 1720, 440), "#F8F0E9", "#B66A32", "云端认知服务层", [
        "中文 ASR 语音识别", "大语言模型与多轮对话", "TTS 语音合成", "Function Call / 知识服务"
    ])

    rounded_box(draw, (260, 690, 720, 970), "#F3EFF8", "#76559B", "数据与可靠性机制", [
        "TLS 鉴权、超时与指数退避", "SD 卡对话日志与语音缓存", "NVS 状态持久化", "弱网检测与离线语音降级"
    ])
    rounded_box(draw, (1080, 690, 1540, 970), "#F7F3E7", "#9B7B29", "交互输出与场景应用", [
        "唤醒—听取—理解—回答动画", "情绪提示与主动陪伴", "灯光等家居控制", "性能、稳定性与用户体验评测"
    ])

    arrow(draw, (500, 295), (690, 295), "采集 / 控制")
    arrow(draw, (1110, 295), (1300, 295), "HTTPS / JSON")
    arrow(draw, (690, 350), (500, 350), "音频 / UI / 指令")
    arrow(draw, (1300, 350), (1110, 350), "文本 / 音频 / 工具调用")
    arrow(draw, (900, 480), (720, 700), "状态与数据")
    arrow(draw, (900, 480), (1080, 700), "交互策略")

    draw.text((900, 1055), "端侧负责实时感知、状态决策与交互执行；云端负责复杂语义理解与生成",
              font=font(22), fill="#455168", anchor="ma")
    image.save(DIAGRAM, quality=95)


REPLACEMENTS = {
    "随着全球人口老龄化加剧与独居青年群体的迅速扩大":
        "随着人口老龄化、家庭小型化以及独居生活方式的发展，面向老年人、独居青年和长期居家人群的情感陪伴需求逐渐受到关注。"
        "传统智能音箱和语音助手主要采用“唤醒—问答”的被动交互方式，对用户是否在场、情绪线索、生活规律和拒绝反馈缺少持续建模，"
        "容易出现交互机械或主动行为不合时宜的问题。因此，研究能够感知环境上下文、控制主动交互边界并提供可靠反馈的智能终端，"
        "对人机交互、情感计算和嵌入式智能系统具有研究意义与应用价值。",
    "本项目以\"Julia\"AI 陪伴终端为研究载体":
        "本项目以“Julia”AI 陪伴终端为研究载体，基于微雪 ESP32-S3 1.85 寸 LCD 开发板，"
        "研究资源受限嵌入式设备上的端云协同智能交互方法。系统以语音、惯性运动、时间与设备状态等信息作为环境上下文，"
        "由端侧完成唤醒检测、状态管理、界面表达、长期记忆与设备控制，云端完成语音识别、大语言模型推理和语音合成。"
        "研究重点包括：构建“感知—决策—表达”闭环架构；设计包含 6 个主状态和 20 个子状态的层次状态机，"
        "实现被动问答、主动关怀与用户拒绝后的节制交互；研究 Function Call 与家居设备抽象层的映射机制。"
        "本研究可为低成本陪伴型智能终端的架构设计、可靠性优化和交互评测提供工程方法与实验依据。",
    "2.2 大语言模型的边缘部署": "2.2 大语言模型的端云协同部署",
    "大语言模型的边缘部署是降低延迟、保护隐私的关键路径":
        "大语言模型的端云协同部署是兼顾交互能力、实时性与隐私保护的重要路径。受限于 ESP32-S3 的存储和算力，"
        "本研究不在端侧直接运行十亿参数级语言模型，而是由端侧承担唤醒、VAD、状态推理、缓存和隐私过滤，"
        "由云端模型承担复杂语义理解与生成。该任务划分能够降低无效云调用，并为弱网降级和本地可控行为保留确定性。",
    "本研究创新性地将情感维度引入层次状态机设计":
        "本研究将情感线索与环境上下文引入层次状态机设计，构建“深度睡眠—待机—陪伴—主动交互—对话—静默”"
        "六级主状态及 20 个子状态，通过事件、超时和用户反馈约束主动行为，探索兼顾主动性与不过度打扰的陪伴节奏。",
    "本项目旨在设计并实现一套面向情感陪伴的端云协同多模态智能交互终端":
        "本项目旨在设计并实现一套面向情感陪伴的端云协同多模态智能交互终端（代号“Julia”），实现以下目标："
        "（1）构建可测量、可降级的中文语音交互链路，在稳定网络下统计端到端时延、识别成功率与失败恢复时间；"
        "（2）设计由情感线索和环境上下文驱动的主动交互状态机，实现从被动响应到适度主动陪伴的转变；"
        "（3）实现大模型 Function Call 与设备控制抽象层的融合；（4）在 ESP32-S3 平台完成系统部署，"
        "通过连续运行、弱网、存储和功耗实验验证系统可靠性。",
    "（2）基于情感计算的层次状态机交互模型":
        "（2）基于情感线索与环境上下文的层次状态机交互模型。构建 6 个主状态（S0—S5）和 20 个子状态的交互空间，"
        "定义用户离开/返回、情绪线索、作息变化、唤醒、低电量与充电等事件。研究状态转换与人物动画、灯效、提示音和语音回复的联动机制。",
    "（4）多协议智能家居控制接口":
        "（4）可扩展的智能家居控制接口。设计统一设备控制抽象层和 JSON 指令模式，首先完成 Wi-Fi/虚拟设备控制闭环，"
        "并为 Matter 与 BLE Mesh 预留适配接口。通过灯光等典型设备验证自然语言、Function Call 与确定性控制指令之间的映射。",
    "（2）情感状态机的最优策略学习问题":
        "（2）主动陪伴策略的可解释决策与边界控制问题。如何融合语音活动、运动、时间、历史交互和用户反馈，"
        "在主动关怀与避免打扰之间取得平衡，并使状态转换可解释、可复现，是交互策略设计的关键问题。",
    "系统采用\"端-边-云\"三层架构":
        "系统采用端云协同架构。ESP32-S3 端负责传感数据采集、本地语音前端、层次状态机、长期记忆、界面渲染、"
        "功耗管理和设备控制；云端负责 ASR、LLM 与 TTS。SD 卡和 NVS 用于对话摘要、画像、状态与离线语音缓存，"
        "网络层提供 TLS、超时重试、信号质量检测和弱网降级。系统总体架构如图 1 所示。",
    "（1）驱动层：基于 ESP-IDF 5.4":
        "（1）驱动层：基于 ESP-IDF 5.5.4 实现 ST77916 LCD、I2S 音频、WS2812/LEDC、ADC 电池检测、"
        "QMI8658 IMU、PCF85063 RTC 与 SD 卡驱动。",
    "（1）功能验证：逐项测试需求文档中的功能点":
        "（1）功能验证：测试唤醒、录音、ASR、LLM、TTS、多轮对话、长期记忆、情绪状态联动和设备控制。"
        "在统一语料和网络条件下统计唤醒成功率、字错误率、端到端时延中位数/P95、失败恢复率及控制成功率，"
        "避免以单次演示结果替代量化评测。",
    "创新点一：提出面向情感陪伴的\"感知-认知-表达\"三层状态机模型":
        "创新点一：提出面向情感陪伴的“感知—决策—表达”层次状态模型。构建 6 主状态、20 子状态的可解释交互空间，"
        "将用户在场、作息、语音情绪线索、拒绝反馈与电源状态统一映射为事件，形成具有主动性边界的陪伴策略。",
    "创新点二：设计端云协同的流式多模态交互架构":
        "创新点二：设计面向资源受限终端的端云协同可靠交互架构。通过端侧唤醒/VAD、状态决策、缓存与隐私过滤，"
        "结合云端 ASR/LLM/TTS，并引入时间同步、网络质量门控、超时重试和离线语音降级机制，提升弱网环境下的可用性。",
    "创新点三：实现大模型 Function Call 与多协议智能家居控制的深度融合":
        "创新点三：设计大模型 Function Call 与确定性设备控制之间的安全映射层。采用结构化 JSON 校验、设备抽象接口和执行结果反馈，"
        "降低自然语言生成的不确定性，并为 Wi-Fi、Matter 和 BLE Mesh 适配提供统一扩展点。",
    "（1）学术成果：完成硕士学位论文 1 篇；在计算机视觉/人机交互领域顶级会议":
        "（1）学术成果：完成硕士学位论文 1 篇；围绕端云协同交互、主动陪伴状态管理或嵌入式系统可靠性形成论文或专利成果，"
        "具体投稿层级根据实验完整性和创新性确定。",
}


REFERENCES = [
    "[1] Picard R W. Affective Computing[M]. Cambridge: MIT Press, 1997.",
    "[2] Calvo R A, D'Mello S. Affect Detection: An Interdisciplinary Review of Models, Methods, and Their Applications[J]. IEEE Transactions on Affective Computing, 2010, 1(1): 18-37.",
    "[3] Poria S, Cambria E, Bajpai R, et al. A Review of Affective Computing: From Unimodal Analysis to Multimodal Fusion[J]. Information Fusion, 2017, 37: 98-125.",
    "[4] Zhang P, Zeng G, Wang T, et al. TinyLlama: An Open-Source Small Language Model[EB/OL]. arXiv:2401.02385, 2024.",
    "[5] Warden P, Situnayake D. TinyML: Machine Learning with TensorFlow Lite on Arduino and Ultra-Low-Power Microcontrollers[M]. O'Reilly Media, 2019.",
    "[6] Espressif Systems. ESP32-S3 Series Datasheet[EB/OL]. 2025.",
    "[7] Espressif Systems. ESP-SR Speech Recognition Framework Programming Guide[EB/OL]. 2025.",
    "[8] OASIS. MQTT Version 5.0[S]. 2019.",
    "[9] Connectivity Standards Alliance. Matter Specification Version 1.3[S]. 2024.",
    "[10] Satyanarayanan M. The Emergence of Edge Computing[J]. Computer, 2017, 50(1): 30-39.",
    "[11] Shi W, Cao J, Zhang Q, et al. Edge Computing: Vision and Challenges[J]. IEEE Internet of Things Journal, 2016, 3(5): 637-646.",
    "[12] Amershi S, Weld D, Vorvoreanu M, et al. Guidelines for Human-AI Interaction[C]//CHI 2019. New York: ACM, 2019: 1-13.",
    "[13] 李沐, 阿斯顿·张, 扎卡里·C. 立顿, 等. 动手学深度学习[M]. 北京: 人民邮电出版社, 2023.",
    "[14] 王飞跃, 等. 具身智能：概念、架构与挑战[J]. 自动化学报, 2024, 50(1): 1-15.",
]


def replace_paragraph(paragraph):
    text = paragraph.text.strip()
    for prefix, replacement in REPLACEMENTS.items():
        if text.startswith(prefix):
            paragraph.text = replacement
            return


def insert_after(paragraph, new_paragraph):
    paragraph._p.addnext(new_paragraph._p)


def revise_document():
    if not BACKUP.exists():
        copy2(SOURCE, BACKUP)
    document = Document(SOURCE)
    for paragraph in document.paragraphs:
        replace_paragraph(paragraph)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_paragraph(paragraph)

    # Replace the original reference list while retaining the section heading.
    start = next((i for i, p in enumerate(document.paragraphs) if p.text.strip() == "八、参考文献"), None)
    if start is not None:
        for p in document.paragraphs[start + 1:]:
            if p.text.strip().startswith("["):
                p._element.getparent().remove(p._element)
        anchor = document.paragraphs[start]
        for reference in reversed(REFERENCES):
            p = document.add_paragraph(reference)
            p.style = document.styles["Normal"]
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            insert_after(anchor, p)

    # Insert the architecture figure after section 4.1 body text.
    target = next((p for p in document.paragraphs if p.text.startswith("系统采用端云协同架构")), None)
    if target is not None:
        pic = document.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.add_run().add_picture(str(DIAGRAM), width=Inches(6.3))
        caption = document.add_paragraph("图 1  Julia 端云协同多模态智能交互系统总体架构")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_after = Pt(6)
        insert_after(target, caption)
        insert_after(target, pic)

    document.core_properties.title = "面向情感陪伴的端云协同多模态智能交互终端设计与实现"
    document.core_properties.subject = "硕士学位论文开题报告（修改定稿版）"
    document.save(OUTPUT)


if __name__ == "__main__":
    make_diagram()
    revise_document()
    print(OUTPUT)
    print(DIAGRAM)
